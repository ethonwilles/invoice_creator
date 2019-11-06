import docx
import os


def creator():
    try:
        inv_num = input("Enter Invoice Number: ")
        location = input("Enter Location: ")
        lot_num = input("Enter Lot Number: ")
        window_amt = input("Enter amount of windows: ")
        window_price_per = 22
        sliders = input("Enter amount of sliders: ")
        slider_price_per = 23
        doors = input("Enter amount of doors: ")
        doors_price_per = 22
        lgWin = input("Enter amount of large windows: ")
        lgWin_price_per = 40
        price_of_trip = input("Enter Price of Trip: ")
        desc = input("Enter description of services if applicable: ")

        total = ((int(window_amt) * int(window_price_per)) + (int(sliders) * int(slider_price_per)) + (int(doors) * int(doors_price_per))+(int(lgWin) * int(lgWin_price_per)) +int(price_of_trip))


        document = docx.Document()

        document.add_heading(f'                        Invoice #{inv_num}', 0)

        p = document.add_paragraph("""
        Chad J Willes Construction
        371 west 500 south
        Lehi UT, 84043
        801-706-8523
        """)


        document.add_paragraph("""
        To:
        Alside Exterior Building Products
        915 West 2610 South 
        Salt Lake City Ut 84119
        Att: Andrew Germaine

        """)


        main_p = document.add_paragraph("")
        main_p.add_run('This invoice is for all the work needed to install windows to the Alside Standard').bold = True
        main_p.add_run(f"\n {location}")
        main_p.add_run(f"\nLot: {lot_num}")
        main_p.add_run(f"\nWindows                   @${window_price_per} per x{window_amt}")
        main_p.add_run(f"\nSliders                        @${slider_price_per} per x{sliders}")
        main_p.add_run(f"\nDoors                          @${doors_price_per} per x{doors}")
        main_p.add_run(f"\nLarge Windows       @${lgWin_price_per} per x{lgWin}")
        main_p.add_run(f"\nTrip Price: ${price_of_trip}")
        main_p.add_run(f"\n{desc}")
        main_p.add_run(f"\nTOTAL                                                                        ${total}").underline = True






        
        try:
            document.save('C:/Users/Owner/Documents/Invoices/'+f'Invoice for Alside {inv_num} {location}.docx')
            creator()
        except:
            print("\nInvoice Already Exists, Try Again.")
            creator()
    except:
        print("/n Program Broke for some reason. Try Again")
        creator()

creator()